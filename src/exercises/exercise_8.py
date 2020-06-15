
# Standard imports:
import os
import numpy as np
import matplotlib.pyplot as plt
from tools.waipy.lib import waipy
from docx.shared import Inches

# Local imports:
from generators import grng, colorednoise, pmodel, logis, henon
from tools import createdocument, getdata


def run(doc_report):
    doc_report.add_heading('Exercise 8', level=2)
    doc_report.add_heading('Exercise 8.1', level=3)
    doc_report.add_paragraph("""
      Here we compare the Continuous Wavelet Spectrum for time series generated with each signal generator used so far,
      along with provided data series. Both Morley and DOG wavelet charts are used.""")

    mount_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), '..', 'mount')
    with open(os.path.join(mount_path, 'surftemp504.txt'), 'r') as text_file:
        list_surftemp504 = [float(w) for w in text_file.read().split('\n')]
        print(len(list_surftemp504))

    with open(os.path.join(mount_path, 'sol3ghz.dat'), 'r') as text_file:
        list_sol3ghz = [float(w) for w in text_file.read().split('\n')]
        print(len(list_sol3ghz))

    # Henon map not used as it causes issues with the waipy module.
    names = ('surftemp504', 'sol3ghz', 'USA_COVID19', 'GNRG', 'Color', 'P_model_038_endogen_beta04',
             'P_model_025_exogen_beta04', 'logistic_rho3.88_tau1.1', 'henon_a1.38_b0.22')
    comments = ('', 'Particularly for sol3ghz data set, we see two wavelet spectrum peaks, which indicates some ' +
                'recurring feature of the signal that occurs in the time associated with around 8000 samples.',
                'The DOG wavelet seems to reveal some interesting pattern on the USA COVID-19 data around the 128 day' +
                'period, while the Morlet transform peaks the spectrum close to 64 days.',
                '', '', '', '', '', '')
    generators = (lambda: list_surftemp504,
                  lambda: list_sol3ghz,
                  lambda: getdata.acquire_data(date_ini='2020-02-20').new_cases.to_list(),
                  lambda: grng.time_series(8192, 1),
                  lambda: colorednoise.powerlaw_psd_gaussian(1, 8192),
                  lambda: pmodel.pmodel(n_values=8192, p=0.38, slope=0.4)[1],
                  lambda: pmodel.pmodel(n_values=8192, p=0.25, slope=0.4)[1],
                  lambda: logis.logistic_series(3.88, 1.1, 8192)[1],
                  lambda: henon.henon_series(np.random.uniform(1.35, 1.4), np.random.uniform(0.21, 0.31), 8192)[1])

    full_names = ('surftemp504 Dataset', 'sol3ghz Dataset', 'Daily new cases of COVID-19 in the USA',
                  'Non Gaussian Random Generator', 'Colored Noise Generator', 'P-Model Endogenous', 'P-Model Exogenous',
                  'Logistic Map', 'Henon Map')
    for (name, func, full_name, comment) in zip(names, generators, full_names, comments):
        data = func()
        doc_report.add_heading(full_name, level=3)

        fig_name = os.path.join(os.path.dirname(os.path.realpath(__file__)), '..', 'mount', name + '_waipy')
        fig_name = os.path.relpath(fig_name, os.getcwd())
        try:
            result = waipy.cwt(data, 1, 1, 0.25, 4, 4 / 0.25, 0.72, 6, mother='Morlet', name='test name')
            waipy.wavelet_plot(fig_name, np.linspace(0, len(data), len(data)), data, 0.03125, result)
            doc_report.add_heading("Morley:", level=4)
            doc_report.document.add_picture(fig_name + '.png', width=Inches(6))
        except Exception as e:
            doc_report.add_heading("Morley could not be computed for " + full_name, level=4)
            doc_report.add_paragraph("The received error message was: \n" + str(e))
        try:
            result = waipy.cwt(data, 1, 1, 0.25, 4, 4 / 0.25, 0.72, 6, mother='DOG', name='test name')
            waipy.wavelet_plot(fig_name, np.linspace(0, len(data), len(data)), data, 0.03125, result)
            doc_report.add_heading("DOG:", level=4)
            doc_report.document.add_picture(fig_name + '.png', width=Inches(6))
        except Exception as e:
            doc_report.add_heading("DOG could not be computed for " + full_name, level=4)
            doc_report.add_paragraph("The received error message was: \n" + str(e))
        plt.close('all')
        doc_report.add_paragraph(comment)


# Sample execution:
if __name__ == '__main__':

    z = np.linspace(0, 2048, 2048)
    x = np.sin(50 * np.pi * z) + 3.5 * np.random.randn(len(z))
    y = np.cos(50 * np.pi * z + np.pi / 4) + 4 * np.random.randn(len(z))

    data_norm = waipy.normalize(x)
    result0 = waipy.cwt(data_norm, 1, 1, 0.25, 4, 4 / 0.25, 0.72, 6, mother='Morlet', name='test name')
    waipy.wavelet_plot('../mount/Sine with noise', z, data_norm, 0.03125, result0)

    data_norm1 = waipy.normalize(y)
    result1 = waipy.cwt(data_norm1, 1, 1, 0.25, 4, 4 / 0.25, 0.72, 6, mother='DOG', name='y')
    waipy.wavelet_plot('../mount/Cosine with noise', z, data_norm1, 0.03125, result1)

    cross_power, coherence, phase_angle = waipy.cross_wavelet(result0['wave'], result1['wave'])
    figname = '../mount/example3.png'
    waipy.plot_cross('../mount/Crosspower sine and cosine', cross_power, phase_angle, z, result0, result1, figname)
    # Initialize report for debugging:
    test_report = createdocument.ReportDocument()
    # Run function:
    run(test_report)
    test_report.finish()
    print("Finished ", __file__)
    plt.show()
