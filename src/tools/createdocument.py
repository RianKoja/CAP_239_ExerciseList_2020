########################################################################################################################
# Defines a class that handles writing results to a word document.
#
#
# Adapted by Rian Koja to publish in a GitHub repository with specified licence.
########################################################################################################################
import os
from docx import Document
from docx.shared import Inches
from pandas.compat import BytesIO
import matplotlib.pyplot as plt


class ReportDocument:
    def __init__(self):
        self.document = Document()
        self.document.add_heading('Exercise List', level=0)
        self.document.add_heading('Rian Koja', level=1)
        self.document.add_heading('CAP 239 Computational Mathematics', level=1)

        self.file_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), '..', 'mount')
        file_version = 0

        while os.path.isfile(os.path.join(self.file_path, "List_RianKoja_v" + str(file_version) + ".docx")):
            file_version += 1

        self.file_name = "List_RianKoja_v" + str(file_version) + ".docx"

    def add_fig(self, wid=6):
        memfile = BytesIO()
        plt.savefig(memfile)
        self.document.add_picture(memfile, width=Inches(wid))
        memfile.close()

    def finish(self):
        self.document.save(os.path.join(self.file_path, self.file_name))
        print("finished word document file.")

    def add_heading(self, text, level=2):
        self.document.add_heading(text, level=level)

    def add_paragraph(self, text):
        self.document.add_paragraph(text)


if __name__ == "__main__":
    # Just test the functions:
    testDoc = ReportDocument()
    testDoc.finish()
